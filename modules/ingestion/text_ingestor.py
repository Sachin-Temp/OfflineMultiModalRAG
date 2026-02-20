"""
modules/ingestion/text_ingestor.py

Handles ingestion of PDF and DOCX files into TextChunk objects.

Pipeline per file:
  1. Detect file type via python-magic (never trust extension)
  2. PDF path:
       a. Try PyMuPDF native text extraction (preserves layout + bboxes)
       b. Fall back to EasyOCR if page has < 100 chars (scanned PDF)
  3. DOCX path:
       a. docx2txt for clean body text
       b. python-docx for tables, headers, structured content
  4. Clean extracted text (strip headers/footers, normalize whitespace)
  5. Chunk into 500-token pieces with 50-token overlap
       - Never cut mid-sentence
       - Every chunk carries full metadata
  6. Return IngestionResult with list of TextChunk objects

No embeddings are computed here — that is the indexing phase's job.
No Milvus or Tantivy writes happen here — that is Phase 4's job.
"""

import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Tuple, Optional

import fitz                          # PyMuPDF
import pdfplumber
import docx2txt
from docx import Document as DocxDocument
import tiktoken
import magic
from loguru import logger

from config.settings import (
    CHUNK_SIZE_TOKENS,
    CHUNK_OVERLAP_TOKENS,
    UPLOAD_DIR,
)
from models.schemas import TextChunk, BBox, Modality, IngestionResult


# ── Tokenizer (shared, load once) ─────────────────────────────────────────
_TOKENIZER = tiktoken.get_encoding("cl100k_base")


def _count_tokens(text: str) -> int:
    return len(_TOKENIZER.encode(text))


def _encode(text: str) -> List[int]:
    return _TOKENIZER.encode(text)


def _decode(tokens: List[int]) -> str:
    return _TOKENIZER.decode(tokens)


# ── Sentence boundary detection ────────────────────────────────────────────
def _split_into_sentences(text: str) -> List[str]:
    """
    Split text into sentences using regex.
    Handles common abbreviations to avoid false splits.
    Returns list of sentence strings (each ends with its punctuation).
    """
    # Python re requires fixed-width lookbehind. We group abbreviations by length (including the dot).
    
    # Length 3 (2 letters + dot): Mr., Ms., Dr., Sr., Jr., vs., pp., no.
    exclude_3 = r"(?<!(?:Mr|Ms|Dr|Sr|Jr|vs|pp|no)\.)"
    
    # Length 4 (3 letters + dot): Mrs., etc., Fig., vol., e.g., i.e.
    exclude_4 = r"(?<!(?:Mrs|etc|Fig|vol|e\.g|i\.e)\.)"
    
    # Length 5 (4 letters + dot): Prof.
    exclude_5 = r"(?<!Prof\.)"
    
    # Pattern: No abbreviation before dot, dot/punct exists, whitespace follows, uppercase letter follows
    pattern = f"{exclude_3}{exclude_4}{exclude_5}(?<=[.!?])\s+(?=[A-Z])"
    
    parts = re.split(pattern, text.strip())
    # Filter empty
    return [p.strip() for p in parts if p.strip()]


# ── Text cleaning ──────────────────────────────────────────────────────────
def _clean_text(text: str) -> str:
    """
    Remove common PDF artifacts, normalize whitespace.
    Does NOT remove content — only cleans formatting noise.
    """
    # Remove form feed characters
    text = text.replace("\f", "\n")
    # Collapse multiple spaces into one
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Collapse more than 2 consecutive newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove hyphenation at line breaks (common in PDFs)
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(lines).strip()


def _remove_headers_footers(pages_text: List[Tuple[int, str]]) -> List[Tuple[int, str]]:
    """
    Heuristically detect and remove repeated text at the top/bottom of pages
    (typical headers/footers). A line is considered a header/footer if it
    appears on more than 60% of pages at the same relative position.

    Args:
        pages_text: list of (page_number, page_text) tuples

    Returns:
        Cleaned list of (page_number, page_text) tuples
    """
    if len(pages_text) < 3:
        return pages_text   # Not enough pages to detect patterns

    threshold = 0.6
    n_pages = len(pages_text)

    # Collect first and last lines per page
    first_lines: List[str] = []
    last_lines: List[str] = []

    for _, text in pages_text:
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        first_lines.append(lines[0] if lines else "")
        last_lines.append(lines[-1] if lines else "")

    # Find repeated first/last lines
    from collections import Counter
    first_counts = Counter(first_lines)
    last_counts  = Counter(last_lines)

    header_candidates = {line for line, count in first_counts.items()
                         if line and count / n_pages >= threshold}
    footer_candidates = {line for line, count in last_counts.items()
                         if line and count / n_pages >= threshold}

    cleaned = []
    for page_num, text in pages_text:
        lines = text.splitlines()
        # Remove header
        if lines and lines[0].strip() in header_candidates:
            lines = lines[1:]
        # Remove footer
        if lines and lines[-1].strip() in footer_candidates:
            lines = lines[:-1]
        cleaned.append((page_num, "\n".join(lines)))

    return cleaned


# ── Chunker ────────────────────────────────────────────────────────────────
def _chunk_text(
    text: str,
    source_file: str,
    page_number: int,
    session_id: Optional[str],
    bbox: Optional[BBox] = None,
    char_offset: int = 0,
) -> List[TextChunk]:
    """
    Split text into overlapping token-bounded chunks.
    Respects sentence boundaries — never cuts mid-sentence.

    Strategy:
      1. Split text into sentences.
      2. Greedily pack sentences into a chunk until CHUNK_SIZE_TOKENS is reached.
      3. When limit hit, save chunk, backtrack by CHUNK_OVERLAP_TOKENS worth of
         sentences to create overlap, continue.
      4. Each chunk records its character offsets for citation purposes.
    """
    sentences = _split_into_sentences(text)
    if not sentences:
        return []

    chunks: List[TextChunk] = []
    current_sentences: List[str] = []
    current_tokens: int = 0
    now = datetime.now(timezone.utc).isoformat()

    i = 0
    while i < len(sentences):
        sentence = sentences[i]
        sentence_tokens = _count_tokens(sentence)

        # Single sentence longer than chunk size — force include it alone
        if sentence_tokens >= CHUNK_SIZE_TOKENS:
            if current_sentences:
                # Save what we have first
                chunk_text = " ".join(current_sentences)
                chunks.append(_make_chunk(
                    chunk_text, source_file, page_number,
                    session_id, bbox, char_offset, now
                ))
                current_sentences, current_tokens = [], 0

            chunks.append(_make_chunk(
                sentence, source_file, page_number,
                session_id, bbox, char_offset, now
            ))
            i += 1
            continue

        if current_tokens + sentence_tokens <= CHUNK_SIZE_TOKENS:
            current_sentences.append(sentence)
            current_tokens += sentence_tokens
            i += 1
        else:
            # Save current chunk
            chunk_text = " ".join(current_sentences)
            chunks.append(_make_chunk(
                chunk_text, source_file, page_number,
                session_id, bbox, char_offset, now
            ))

            # Build overlap: walk back until we've included CHUNK_OVERLAP_TOKENS
            overlap_sentences: List[str] = []
            overlap_tokens: int = 0
            for sent in reversed(current_sentences):
                t = _count_tokens(sent)
                if overlap_tokens + t > CHUNK_OVERLAP_TOKENS:
                    break
                overlap_sentences.insert(0, sent)
                overlap_tokens += t

            current_sentences = overlap_sentences
            current_tokens = overlap_tokens
            # Do NOT increment i — re-evaluate current sentence with new context

    # Flush remainder
    if current_sentences:
        chunk_text = " ".join(current_sentences)
        chunks.append(_make_chunk(
            chunk_text, source_file, page_number,
            session_id, bbox, char_offset, now
        ))

    return chunks


def _make_chunk(
    text: str,
    source_file: str,
    page_number: int,
    session_id: Optional[str],
    bbox: Optional[BBox],
    char_offset: int,
    timestamp: str,
) -> TextChunk:
    return TextChunk(
        chunk_id=str(uuid.uuid4()),
        text=text,
        modality=Modality.TEXT,
        source_file=source_file,
        page_number=page_number,
        bbox=bbox,
        char_start=char_offset,
        char_end=char_offset + len(text),
        ingest_timestamp=timestamp,
        session_id=session_id,
    )


# ── PDF Extraction ─────────────────────────────────────────────────────────
def _is_scanned_page(page_text: str) -> bool:
    """Return True if page has fewer than 100 characters (likely scanned)."""
    return len(page_text.strip()) < 100


def _extract_pdf_page_native(page: fitz.Page) -> Tuple[str, Optional[BBox]]:
    """
    Extract text from a single PDF page using PyMuPDF.
    Also extracts bounding box of the text block for citation use.
    Returns (text, bbox) tuple.
    """
    blocks = page.get_text("dict")["blocks"]
    page_text_parts = []
    x0_vals, y0_vals, x1_vals, y1_vals = [], [], [], []

    for block in blocks:
        if block.get("type") == 0:  # type 0 = text block
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    span_text = span.get("text", "").strip()
                    if span_text:
                        page_text_parts.append(span_text)
                        bbox = span.get("bbox", None)
                        if bbox:
                            x0_vals.append(bbox[0])
                            y0_vals.append(bbox[1])
                            x1_vals.append(bbox[2])
                            y1_vals.append(bbox[3])

    full_text = " ".join(page_text_parts)
    combined_bbox = None
    if x0_vals:
        combined_bbox = BBox(
            x0=min(x0_vals), y0=min(y0_vals),
            x1=max(x1_vals), y1=max(y1_vals)
        )
    return full_text, combined_bbox


def _extract_pdf_page_ocr(image_bytes: bytes) -> str:
    """
    Fallback: run EasyOCR on a page rendered as an image.
    Used when native PyMuPDF extraction yields < 100 chars (scanned PDF).
    Always runs on CPU to preserve VRAM for LLM.
    """
    import easyocr
    import numpy as np
    from PIL import Image
    import io

    logger.debug("OCR fallback triggered for scanned page")
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    img_array = np.array(img)

    # EasyOCR on CPU — gpu=False is mandatory here
    reader = easyocr.Reader(["en"], gpu=False, verbose=False)
    results = reader.readtext(img_array, detail=0, paragraph=True)
    return " ".join(results)


def _ingest_pdf(
    file_path: Path,
    session_id: Optional[str],
) -> Tuple[List[TextChunk], List[str]]:
    """
    Full PDF ingestion pipeline.
    Returns (list of TextChunks, list of warning strings).
    """
    chunks: List[TextChunk] = []
    warnings: List[str] = []
    source_file = file_path.name

    try:
        doc = fitz.open(str(file_path))
    except Exception as e:
        return [], [f"PyMuPDF failed to open {file_path.name}: {e}"]

    pages_raw: List[Tuple[int, str, Optional[BBox]]] = []

    for page_idx in range(len(doc)):
        page = doc[page_idx]
        page_num = page_idx + 1    # 1-indexed for citations

        native_text, bbox = _extract_pdf_page_native(page)

        if _is_scanned_page(native_text):
            logger.info(f"Page {page_num} of {source_file} is scanned — using OCR fallback")
            warnings.append(f"Page {page_num}: OCR fallback used (scanned/low-text page)")
            try:
                pix = page.get_pixmap(dpi=200)
                image_bytes = pix.tobytes("png")
                ocr_text = _extract_pdf_page_ocr(image_bytes)
                pages_raw.append((page_num, ocr_text, None))
            except Exception as e:
                warnings.append(f"Page {page_num}: OCR also failed — {e}")
                logger.warning(f"OCR failed on page {page_num}: {e}")
        else:
            pages_raw.append((page_num, native_text, bbox))

    doc.close()

    # Remove headers/footers
    pages_text_only = [(pn, txt) for pn, txt, _ in pages_raw]
    cleaned_pages = _remove_headers_footers(pages_text_only)

    # Re-attach bboxes after header/footer removal
    bbox_map = {pn: b for pn, _, b in pages_raw}

    for page_num, page_text in cleaned_pages:
        page_text = _clean_text(page_text)
        if not page_text.strip():
            continue
        bbox = bbox_map.get(page_num)
        page_chunks = _chunk_text(
            text=page_text,
            source_file=source_file,
            page_number=page_num,
            session_id=session_id,
            bbox=bbox,
            char_offset=0,
        )
        chunks.extend(page_chunks)

    logger.info(f"PDF ingestion complete: {source_file} → {len(chunks)} text chunks")
    return chunks, warnings


# ── DOCX Extraction ────────────────────────────────────────────────────────
def _extract_docx_tables(doc_path: Path) -> str:
    """
    Extract text from DOCX tables using python-docx.
    Tables are often missed by docx2txt.
    Returns table content as pipe-delimited plain text.
    """
    doc = DocxDocument(str(doc_path))
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                table_texts.append(row_text)
    return "\n".join(table_texts)


def _extract_docx_structured(doc_path: Path) -> str:
    """
    Extract text from DOCX preserving heading hierarchy.
    Returns structured text with heading markers.
    """
    doc = DocxDocument(str(doc_path))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Mark headings so chunker can use them as natural break points
        if para.style.name.startswith("Heading"):
            parts.append(f"\n## {text}\n")
        else:
            parts.append(text)
    return "\n".join(parts)


def _ingest_docx(
    file_path: Path,
    session_id: Optional[str],
) -> Tuple[List[TextChunk], List[str]]:
    """
    Full DOCX ingestion pipeline.
    Returns (list of TextChunks, list of warning strings).
    DOCX files don't have pages — all chunks get page_number=1.
    """
    chunks: List[TextChunk] = []
    warnings: List[str] = []
    source_file = file_path.name

    # Extract body text
    try:
        body_text = docx2txt.process(str(file_path))
    except Exception as e:
        warnings.append(f"docx2txt failed: {e} — falling back to python-docx")
        try:
            body_text = _extract_docx_structured(file_path)
        except Exception as e2:
            return [], [f"DOCX extraction completely failed: {e2}"]

    # Extract tables separately and append
    try:
        table_text = _extract_docx_tables(file_path)
        if table_text.strip():
            body_text += "\n\n" + table_text
    except Exception as e:
        warnings.append(f"Table extraction failed: {e}")

    body_text = _clean_text(body_text)

    if not body_text.strip():
        return [], [f"No text could be extracted from {source_file}"]

    # DOCX has no pages — treat entire document as page 1
    page_chunks = _chunk_text(
        text=body_text,
        source_file=source_file,
        page_number=1,
        session_id=session_id,
        bbox=None,
        char_offset=0,
    )
    chunks.extend(page_chunks)

    logger.info(f"DOCX ingestion complete: {source_file} → {len(chunks)} text chunks")
    return chunks, warnings


# ── File Type Detection ────────────────────────────────────────────────────
def _detect_mime_type(file_path: Path) -> str:
    """
    Detect file MIME type using libmagic (never trust extension).
    Returns MIME type string, e.g. 'application/pdf'.
    """
    try:
        mime = magic.from_file(str(file_path), mime=True)
        return mime
    except Exception as e:
        logger.warning(f"python-magic failed ({e}), falling back to extension detection")
        suffix = file_path.suffix.lower()
        fallback = {
            ".pdf":  "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc":  "application/msword",
        }
        return fallback.get(suffix, "application/octet-stream")


# ── Public API ─────────────────────────────────────────────────────────────
class TextIngestor:
    """
    Public interface for text ingestion.
    Usage:
        ingestor = TextIngestor()
        result = ingestor.ingest(Path("path/to/file.pdf"), session_id="abc123")
        # result.text_chunks contains list of TextChunk objects
    """

    SUPPORTED_MIME_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "docx",
    }

    def ingest(
        self,
        file_path: Path,
        session_id: Optional[str] = None,
    ) -> IngestionResult:
        """
        Ingest a single PDF or DOCX file.

        Args:
            file_path:  Absolute or relative path to the file.
            session_id: Optional session identifier for multi-session support.

        Returns:
            IngestionResult with text_chunks populated and any errors/warnings recorded.
        """
        file_path = Path(file_path).resolve()
        result = IngestionResult(
            source_file=file_path.name,
            modality=Modality.TEXT,
        )

        if not file_path.exists():
            result.errors.append(f"File not found: {file_path}")
            logger.error(f"TextIngestor: File not found: {file_path}")
            return result

        mime_type = _detect_mime_type(file_path)
        file_format = self.SUPPORTED_MIME_TYPES.get(mime_type)

        if file_format is None:
            # Fallback for some windows mime types that magic might miss or classify generically
            # if magic returns "application/x-zip-compressed" for docx, check extension
            if file_path.suffix.lower() in [".docx", ".doc"]:
                 file_format = "docx"
            elif file_path.suffix.lower() == ".pdf":
                 file_format = "pdf"
            else:
                result.errors.append(
                    f"Unsupported file type '{mime_type}' for file {file_path.name}. "
                    f"Supported: PDF, DOCX."
                )
                logger.error(f"TextIngestor: Unsupported MIME type {mime_type}")
                return result

        logger.info(f"TextIngestor: Ingesting {file_path.name} (type={file_format}, session={session_id})")

        try:
            if file_format == "pdf":
                chunks, warnings = _ingest_pdf(file_path, session_id)
            else:
                chunks, warnings = _ingest_docx(file_path, session_id)

            result.text_chunks = chunks
            result.warnings = warnings

            if not chunks:
                result.errors.append(f"No text chunks produced from {file_path.name}")

        except Exception as e:
            result.errors.append(f"Unexpected error during ingestion: {e}")
            logger.exception(f"TextIngestor: Unexpected error on {file_path.name}")

        logger.info(
            f"TextIngestor: {file_path.name} complete — "
            f"{len(result.text_chunks)} chunks, "
            f"{len(result.warnings)} warnings, "
            f"{len(result.errors)} errors"
        )
        return result

    def ingest_batch(
        self,
        file_paths: List[Path],
        session_id: Optional[str] = None,
    ) -> List[IngestionResult]:
        """
        Ingest multiple files sequentially.
        Returns list of IngestionResult, one per file.
        """
        results = []
        for fp in file_paths:
            results.append(self.ingest(fp, session_id))
        return results
