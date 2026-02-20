"""
tests/test_text_ingestor.py
Full test suite for the TextIngestor module.

Tests cover:
- Schema correctness
- Chunker token limits
- Sentence boundary respect
- PDF ingestion (native text)
- DOCX ingestion
- Unsupported file type rejection
- Missing file handling
- Metadata completeness

Run with: pytest tests/test_text_ingestor.py -v
"""

import pytest
import tempfile
from pathlib import Path


# ── Schema tests ───────────────────────────────────────────────────────────
class TestSchemas:

    def test_text_chunk_defaults(self):
        from models.schemas import TextChunk
        chunk = TextChunk(text="hello", source_file="test.pdf", page_number=1)
        assert chunk.chunk_id is not None
        assert len(chunk.chunk_id) == 36   # UUID4 format
        assert chunk.modality.value == "text"
        assert chunk.embedding is None

    def test_text_chunk_to_dict(self):
        from models.schemas import TextChunk, BBox
        bbox = BBox(x0=0, y0=0, x1=100, y1=50)
        chunk = TextChunk(
            text="hello world",
            source_file="test.pdf",
            page_number=2,
            bbox=bbox,
            session_id="sess-001"
        )
        d = chunk.to_dict()
        assert d["text"] == "hello world"
        assert d["page_number"] == 2
        assert d["bbox"]["x0"] == 0
        assert d["session_id"] == "sess-001"
        assert d["modality"] == "text"

    def test_ingestion_result_properties(self):
        from models.schemas import IngestionResult, TextChunk, Modality
        result = IngestionResult(source_file="test.pdf", modality=Modality.TEXT)
        assert result.total_chunks == 0
        assert result.success is False

        chunk = TextChunk(text="hi", source_file="test.pdf", page_number=1)
        result.text_chunks.append(chunk)
        assert result.total_chunks == 1
        assert result.success is True

    def test_image_chunk_defaults(self):
        from models.schemas import ImageChunk
        chunk = ImageChunk(
            image_path="media/img.png",
            source_file="test.pdf",
            page_number=1
        )
        assert chunk.modality.value == "image"
        assert chunk.embedding is None

    def test_audio_chunk_defaults(self):
        from models.schemas import AudioChunk
        chunk = AudioChunk(
            text="hello", audio_file="call.mp3",
            start_time=0.0, end_time=30.0
        )
        assert chunk.modality.value == "audio"
        assert chunk.word_timestamps == []


# ── Chunker unit tests ─────────────────────────────────────────────────────
class TestChunker:

    def _run_chunker(self, text):
        from modules.ingestion.text_ingestor import _chunk_text
        return _chunk_text(
            text=text,
            source_file="test.pdf",
            page_number=1,
            session_id="test-session",
        )

    def test_short_text_single_chunk(self):
        chunks = self._run_chunker("Hello world. This is a short sentence.")
        assert len(chunks) == 1
        assert "Hello world" in chunks[0].text

    def test_chunk_token_limit(self):
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        from config.settings import CHUNK_SIZE_TOKENS

        # Generate text that definitely exceeds one chunk
        sentence = "The quick brown fox jumps over the lazy dog. "
        long_text = sentence * 100   # ~900+ tokens

        chunks = self._run_chunker(long_text)
        for chunk in chunks:
            token_count = len(enc.encode(chunk.text))
            # Allow slight overage for single long sentences
            assert token_count <= CHUNK_SIZE_TOKENS + 50, \
                f"Chunk exceeded token limit: {token_count} tokens"

    def test_overlap_exists(self):
        """Adjacent chunks should share some content (overlap)."""
        sentence = "This is sentence number {}. "
        long_text = "".join(sentence.format(i) for i in range(200))
        chunks = self._run_chunker(long_text)

        if len(chunks) > 1:
            # Last words of chunk N should appear in first part of chunk N+1
            end_of_first = chunks[0].text.split()[-5:]
            start_of_second = chunks[1].text.split()[:20]
            overlap = set(end_of_first) & set(start_of_second)
            assert len(overlap) > 0, "No overlap detected between adjacent chunks"

    def test_metadata_completeness(self):
        chunks = self._run_chunker("Hello. This is a test sentence for metadata checking.")
        for chunk in chunks:
            assert chunk.chunk_id is not None
            assert chunk.source_file == "test.pdf"
            assert chunk.page_number == 1
            assert chunk.session_id == "test-session"
            assert chunk.ingest_timestamp is not None
            assert chunk.modality.value == "text"

    def test_empty_text_returns_no_chunks(self):
        chunks = self._run_chunker("   \n\n   ")
        assert chunks == []


# ── TextIngestor integration tests ────────────────────────────────────────
class TestTextIngestor:

    @pytest.fixture
    def ingestor(self):
        from modules.ingestion.text_ingestor import TextIngestor
        return TextIngestor()

    @pytest.fixture
    def sample_pdf(self, tmp_path):
        """Create a minimal real PDF using PyMuPDF."""
        import fitz
        pdf_path = tmp_path / "test_doc.pdf"
        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text(
                (72, 100),
                f"Page {i+1} content. This is test paragraph {i+1}. "
                f"The quick brown fox jumps over the lazy dog. "
                f"Machine learning models process natural language. "
                f"Retrieval augmented generation improves accuracy. " * 3,
                fontsize=11
            )
        doc.save(str(pdf_path))
        doc.close()
        return pdf_path

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a minimal DOCX file."""
        from docx import Document
        docx_path = tmp_path / "test_doc.docx"
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph(
            "This is the first paragraph. It contains several sentences. "
            "The quick brown fox jumps over the lazy dog. "
            "Natural language processing enables understanding of text. " * 5
        )
        doc.add_heading("Section 2", 1)
        doc.add_paragraph(
            "Second section content here. More sentences follow. "
            "Information retrieval systems find relevant documents. " * 5
        )
        doc.save(str(docx_path))
        return docx_path

    def test_pdf_ingestion_returns_chunks(self, ingestor, sample_pdf):
        result = ingestor.ingest(sample_pdf, session_id="test-sess")
        assert result.success
        assert len(result.text_chunks) >= 1
        assert result.errors == []

    def test_pdf_chunk_metadata(self, ingestor, sample_pdf):
        result = ingestor.ingest(sample_pdf, session_id="test-sess")
        for chunk in result.text_chunks:
            assert chunk.source_file == sample_pdf.name
            assert chunk.page_number >= 1
            assert chunk.session_id == "test-sess"
            assert chunk.text.strip() != ""
            assert chunk.chunk_id is not None

    def test_docx_ingestion_returns_chunks(self, ingestor, sample_docx):
        result = ingestor.ingest(sample_docx, session_id="test-sess")
        assert result.success
        assert len(result.text_chunks) >= 1

    def test_docx_page_number_is_one(self, ingestor, sample_docx):
        """DOCX has no pages — all chunks should have page_number=1."""
        result = ingestor.ingest(sample_docx, session_id="test-sess")
        for chunk in result.text_chunks:
            assert chunk.page_number == 1

    def test_missing_file_returns_error(self, ingestor, tmp_path):
        fake_path = tmp_path / "nonexistent.pdf"
        result = ingestor.ingest(fake_path)
        assert not result.success
        assert len(result.errors) > 0
        assert "not found" in result.errors[0].lower()

    def test_unsupported_file_type_returns_error(self, ingestor, tmp_path):
        txt_file = tmp_path / "file.txt"
        txt_file.write_text("some content", encoding="utf-8")
        result = ingestor.ingest(txt_file)
        assert not result.success
        # It detects text/plain from magic, so it returns unsupported error
        assert len(result.errors) > 0

    def test_batch_ingest(self, ingestor, sample_pdf, sample_docx):
        results = ingestor.ingest_batch(
            [sample_pdf, sample_docx],
            session_id="batch-test"
        )
        assert len(results) == 2
        assert all(r.success for r in results)

    def test_chunk_text_is_non_empty(self, ingestor, sample_pdf):
        result = ingestor.ingest(sample_pdf)
        for chunk in result.text_chunks:
            assert len(chunk.text.strip()) > 0

    def test_no_duplicate_chunk_ids(self, ingestor, sample_pdf):
        result = ingestor.ingest(sample_pdf)
        ids = [c.chunk_id for c in result.text_chunks]
        assert len(ids) == len(set(ids)), "Duplicate chunk IDs found"
